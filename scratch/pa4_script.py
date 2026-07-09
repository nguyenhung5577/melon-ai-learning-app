import os
import re
import shutil
from docx import Document

base_dir = r"d:\HCMUS\HK2-N3\SE4AI\melon-ai-learning-app"
pa4_dir = os.path.join(base_dir, "pa", "pa4")
os.makedirs(pa4_dir, exist_ok=True)

# 1. Update TEST_PLAN_MELON_AI.md and save as Test cases.md
tp_src = os.path.join(base_dir, "TEST_PLAN_MELON_AI.md")
with open(tp_src, 'r', encoding='utf-8') as f:
    tp_content = f.read()

# Group into 5 Use Cases to satisfy requirements (at least 5 tests per Use Case, up to 10 Use cases)
tp_content = tp_content.replace('### 2.1. Phân hệ Xác thực & Phân quyền (ML-AUTH)', '### 2.1. Use Case 1: Authentication & Security')
tp_content = tp_content.replace('### 2.2. Phân hệ Subscription & Thanh toán Stripe (ML-SUBS)', '### 2.2. Use Case 2: Application Features & Payments\n')
tp_content = re.sub(r'### 2\.(7|11|14)\. Phân hệ.*?\n', '', tp_content)
tp_content = tp_content.replace('### 2.3. Phân hệ OCR & Bóc tách đề thông minh (ML-PARSE)', '### 2.3. Use Case 3: Knowledge Base & Parsing\n')
tp_content = re.sub(r'### 2\.10\. Phân hệ.*?\n', '', tp_content)
tp_content = tp_content.replace('### 2.4. Phân hệ Cá nhân hóa Đề thi (ML-PERSONALIZE)', '### 2.4. Use Case 4: Practice & Personalization\n')
tp_content = re.sub(r'### 2\.5\. Phân hệ.*?\n', '', tp_content)
tp_content = tp_content.replace('### 2.6. Phân hệ Gamification & Xếp hạng (ML-GAMI)', '### 2.5. Use Case 5: Gamification & System Administration\n')
tp_content = re.sub(r'### 2\.8\. Phân hệ.*?\n', '', tp_content)

tc_dest = os.path.join(pa4_dir, "Test cases.md")
with open(tc_dest, 'w', encoding='utf-8') as f:
    f.write(tp_content)

# Calculate dynamic stats for the 5 Use Cases
use_cases = re.split(r'### 2\.\d+\. Use Case', tp_content)
summary_lines = []

for i, uc_text in enumerate(use_cases[1:], 1):
    title = uc_text.split('\n')[0].strip()
    title = title.split(': ', 1)[-1]
    
    x_count = len(re.findall(r'\[x\]', uc_text, re.IGNORECASE))
    f_count = len(re.findall(r'\[!\]', uc_text))
    s_count = len(re.findall(r'\[-\]', uc_text))
    e_count = len(re.findall(r'\[ \]', uc_text))
    total = x_count + f_count + s_count + e_count
    
    summary_lines.append(f"{i}. **{title}:** {total} tests ({x_count} Pass, {f_count} Fail, {s_count + e_count} Skip/Unimplemented)")

total_tests = sum([int(re.search(r'\*\* (\d+) tests', s).group(1)) for s in summary_lines])
total_pass = sum([int(re.search(r'\((\d+) Pass', s).group(1)) for s in summary_lines])
total_fail = sum([int(re.search(r'Pass, (\d+) Fail', s).group(1)) for s in summary_lines])
total_skip = total_tests - total_pass - total_fail

# 2. Update BUG_REPORT_MELON_AI.md and save as Test report.md
br_src = os.path.join(base_dir, "BUG_REPORT_MELON_AI.md")
with open(br_src, 'r', encoding='utf-8') as f:
    br_content = f.read()

summary_report = f"""
## Test Summary Report
- **Number of Use-Cases tested:** 5
- **Total Test Cases:** {total_tests}
- **Passed Test Cases:** {total_pass}
- **Failed Test Cases:** {total_fail}
- **Skipped/Unimplemented Test Cases:** {total_skip}

### Breakdown by Use-Case:
""" + '\n'.join(summary_lines) + "\n\n---\n\n"

br_content = summary_report + br_content

tr_dest = os.path.join(pa4_dir, "Test report.md")
with open(tr_dest, 'w', encoding='utf-8') as f:
    f.write(br_content)

# 3. Fill out rup_tstpln.docx and save as Test plan.docx
doc_src = os.path.join(base_dir, "rup_tstpln.docx")
try:
    doc = Document(doc_src)
    
    for p in doc.paragraphs:
        if "<Project Name>" in p.text:
            p.text = p.text.replace("<Project Name>", "Melon AI Learning App")
        if "<1.0>" in p.text:
            p.text = p.text.replace("<1.0>", "1.0")
        if "[Provide an overview of the test plan document here" in p.text:
            p.text = "This document outlines the test plan for the Melon AI Learning App, an AI-enabled educational platform. Its objective is to define the testing scope, approach, resources, and schedule. It will be used by the QA team, developers, and project managers to ensure software quality."
        if "[List software items for testing." in p.text:
            p.text = "Target Test Items include 5 core use cases:\n1. Authentication & Security\n2. Application Features & Payments\n3. Knowledge Base & Parsing\n4. Practice & Personalization\n5. Gamification & System Administration"
        if "[This section presents the non-human resources required" in p.text:
            p.text = "Test environment requires local machines (PC/Laptop) running Windows, macOS, or Linux, capable of running Next.js and FastAPI backends, and modern web browsers (Chrome, Firefox, Edge)."
        if "[Identify hardware requirements used to run and test the application]" in p.text:
            p.text = "Minimum Hardware: 4-core CPU, 8GB RAM, 50GB storage. Network connection for Stripe and LLM APIs."
    
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "<Project Name>" in p.text:
                        p.text = p.text.replace("<Project Name>", "Melon AI Learning App")
                    if "<1.0>" in p.text:
                        p.text = p.text.replace("<1.0>", "1.0")
    
    doc_dest = os.path.join(pa4_dir, "Test plan.docx")
    doc.save(doc_dest)
except Exception as e:
    print(f"Error processing docx: {e}")
