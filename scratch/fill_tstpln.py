import os
from docx import Document

base_dir = r"d:\HCMUS\HK2-N3\SE4AI\melon-ai-learning-app"
src_doc = os.path.join(base_dir, "rup_tstpln.docx")
dest_doc = os.path.join(base_dir, "pa", "pa4", "Test plan_FINAL.docx")

doc = Document(src_doc)

def replace_in_p(p, old, new):
    if old in p.text:
        p.text = p.text.replace(old, new)

for p in doc.paragraphs:
    replace_in_p(p, "<Project Name>", "Melon App")
    replace_in_p(p, "<1.0>", "2.0")
    
    if "[Provide an overview of the test plan document here" in p.text:
        p.text = """This Test Plan document details the strategies, resources, and schedules for the testing phase of Melon App, an AI-enabled interactive educational platform for children aged 6-12. 
Objectives of this document:
- Define the scope of testing including core functionalities (Authentication, Payments, RAG Parsing, AI Practice, Gamification).
- Identify the required hardware, software, and personnel resources.
- Outline the roles of the testing team members (Vũ, Hoàng, Hùng).
This document should be used by the QA team, developers, and project stakeholders to ensure the final application meets the safety, performance (<2s response), and functional requirements outlined in the Vision Document and SDP."""
    
    if "[List software items for testing. Software items include features and areas to test such as performance, usability, and security.]" in p.text:
        p.text = """Target Test Items:
1. Authentication & Security (ML-AUTH): Firebase Auth integration, role-based access control (Kid vs Parent vs Admin).
2. Application Features & Payments (ML-SUBS, ML-FAM, ML-PROFILE): Stripe checkout, Parent dashboard, child account management.
3. Knowledge Base & Parsing (ML-PARSE, ML-RAG-INGEST): PDF/Image upload, AI-driven math parsing via FastAPI and OpenRouter.
4. Practice & Personalization (ML-PRAC, ML-PERSONALIZE): Interactive quizzes, LLM chat integration, personalized weaknesses tracking.
5. Gamification & System Administration (ML-GAMI, ML-ADM, ML-ERR): EXP progression, badges, leaderboard, robust error handling."""

    if "[This section presents the non-human resources required" in p.text:
        p.text = "This section details the hardware, software, and tools required to effectively test the Melon App Next.js frontend and FastAPI backend."

    if "[Identify hardware requirements used to run and test the application]" in p.text:
        p.text = "Hardware Requirements:\n- Client Testing: Laptops/PCs (Windows 10+, macOS) and mobile devices (iOS/Android tablets) to simulate child interactions.\n- Server Environment: Vercel cloud environment for frontend staging. Local Python servers (FastAPI) requiring min. 8GB RAM for RAG/LLM proxy processing.\n- Network: High-speed internet required to interact with Firebase Storage, OpenAI/OpenRouter APIs, and Stripe Webhooks."

    if "[Note:  Add or delete items as appropriate.]" in p.text:
        p.text = ""
    if "[Note:  Add more items as appropriate.]" in p.text:
        p.text = ""

if len(doc.tables) > 0:
    rev_table = doc.tables[0]
    if len(rev_table.rows) > 1:
        rev_table.rows[1].cells[0].text = "10/Jun/2026"
        rev_table.rows[1].cells[1].text = "2.0"
        rev_table.rows[1].cells[2].text = "Initial Test Plan based on SDP v2"
        rev_table.rows[1].cells[3].text = "Vũ, Hoàng, Hùng"

if len(doc.tables) > 2:
    prod_table = doc.tables[2]
    if len(prod_table.rows) > 1:
        prod_table.rows[1].cells[0].text = "Project Management & Bug Tracking"
        prod_table.rows[1].cells[1].text = "Jira"
        prod_table.rows[1].cells[2].text = "Atlassian"
        prod_table.rows[1].cells[3].text = "Cloud"
    if len(prod_table.rows) > 2:
        prod_table.rows[2].cells[0].text = "Version Control"
        prod_table.rows[2].cells[1].text = "GitHub"
        prod_table.rows[2].cells[2].text = "GitHub"
        prod_table.rows[2].cells[3].text = "Cloud"

if len(doc.tables) > 3:
    roles_table = doc.tables[3]
    if len(roles_table.rows) > 1:
        roles_table.rows[1].cells[1].text = "Hoàng (1)"
    if len(roles_table.rows) > 2:
        roles_table.rows[2].cells[1].text = "Vũ, Hùng (2)"
    if len(roles_table.rows) > 3:
        roles_table.rows[3].cells[1].text = "Hoàng (1)"

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_p(p, "<Project Name>", "Melon App")
                replace_in_p(p, "<1.0>", "2.0")

try:
    doc.save(dest_doc)
except Exception as e:
    print(e)
