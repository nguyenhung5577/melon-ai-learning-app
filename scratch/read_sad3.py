import sys
from docx import Document
sys.stdout.reconfigure(encoding='utf-8')

try:
    doc = Document(r'd:\HCMUS\HK2-N3\SE4AI\melon-ai-learning-app\pa\pa3\SAD-v3.docx')
    for p in doc.paragraphs[:50]:
        if p.text.strip():
            print(p.text)
    print("---TABLES---")
    for t in doc.tables:
        for row in t.rows[:5]:
            print(' | '.join([str(c.text).replace('\n', ' ') for c in row.cells]))
        break
except Exception as e:
    print(e)
