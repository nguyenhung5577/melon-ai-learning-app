import sys
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn

def fix_format():
    path = r'd:\HCMUS\HK2-N3\SE4AI\melon-ai-learning-app\pa\pa5\Test automation report.docx'
    out_path = r'd:\HCMUS\HK2-N3\SE4AI\melon-ai-learning-app\pa\pa5\Test automation report_FIXED.docx'
    
    doc = Document(path)
    
    # Process paragraphs
    for para in doc.paragraphs:
        # Remove paragraph shading (background color)
        pPr = para._p.get_or_add_pPr()
        shd = pPr.find(qn('w:shd'))
        if shd is not None:
            pPr.remove(shd)
            
        for run in para.runs:
            # Remove text highlight
            if run.font.highlight_color is not None:
                run.font.highlight_color = None
                
            # Set font color to auto (or remove explicit color)
            # Setting it to None resets it to default
            run.font.color.rgb = None
            
            # Also remove run background (shading inside run)
            rPr = run._r.get_or_add_rPr()
            shd_r = rPr.find(qn('w:shd'))
            if shd_r is not None:
                rPr.remove(shd_r)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Remove cell shading
                tcPr = cell._tc.get_or_add_tcPr()
                shd_c = tcPr.find(qn('w:shd'))
                if shd_c is not None:
                    tcPr.remove(shd_c)
                    
                for para in cell.paragraphs:
                    pPr = para._p.get_or_add_pPr()
                    shd = pPr.find(qn('w:shd'))
                    if shd is not None:
                        pPr.remove(shd)
                        
                    for run in para.runs:
                        if run.font.highlight_color is not None:
                            run.font.highlight_color = None
                        run.font.color.rgb = None
                        rPr = run._r.get_or_add_rPr()
                        shd_r = rPr.find(qn('w:shd'))
                        if shd_r is not None:
                            rPr.remove(shd_r)

    doc.save(out_path)
    print(f"Saved fixed document to {out_path}")

if __name__ == '__main__':
    fix_format()
