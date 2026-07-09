import os
from docx import Document

base_dir = r"d:\HCMUS\HK2-N3\SE4AI\melon-ai-learning-app"
src_doc = os.path.join(base_dir, "rup_tstpln.docx")
dest_doc = os.path.join(base_dir, "pa", "pa4", "Test plan.docx")
final_doc = os.path.join(base_dir, "pa", "pa4", "Test plan_FINAL.docx")

doc = Document(src_doc)

def replace_in_p(p, old, new):
    if old in p.text:
        p.text = p.text.replace(old, new)

for p in doc.paragraphs:
    replace_in_p(p, "<Project Name>", "Melon Learning App")
    replace_in_p(p, "<1.0>", "3.0")
    
    if "[Provide an overview of the test plan document here" in p.text:
        p.text = """This Test Plan document details the strategies, resources, and schedules for the testing phase of Melon, an AI-assisted educational platform for children, parents, and administrators.
Objectives of this document:
- Define the scope of testing including the Next.js frontend, secure API proxy, and the FastAPI AI backend.
- Evaluate the Hybrid Split-Stack Architecture, AI integrations (GPT-4o-mini, Gemini 2.0 Flash, ElevenLabs TTS), and Vector Search (Pinecone).
- Identify required hardware, software, and personnel resources.
- Outline the roles of the testing team members (Vũ, Hoàng, Hùng).
This document guides the QA team and developers to ensure the system meets functional constraints and child-safety requirements as specified in the SAD v3."""
    
    if "[List software items for testing." in p.text:
        p.text = """Target Test Items:
1. Authentication & Security (ML-AUTH): Firebase Auth, Role-based Access, Child-safety moderation filters.
2. Application Features & Payments (ML-SUBS, ML-PROFILE, ML-FAM): Parent Dashboard, Stripe Subscriptions, State management (Zustand).
3. Knowledge Base & Parsing (ML-PARSE, ML-RAG-INGEST): Multimodal problem parsing (Gemini 2.0), OCR pipelines, PDF extraction via Python.
4. Practice & Personalization (ML-PRAC, ML-PERSONALIZE): RAG-based Quiz Generation (GPT-4o-mini), Speech-to-Text (Whisper), TTS (ElevenLabs).
5. Gamification & Administration (ML-GAMI, ML-ADM): Progress tracking, Leaderboards, Admin Question Bank management."""

    if "[This section presents the non-human resources required" in p.text:
        p.text = "This section details the hardware, software, and tools required to test the Melon Hybrid Architecture (Next.js 16 + FastAPI)."

    if "[Identify hardware requirements used to run and test the application]" in p.text:
        p.text = "Hardware Requirements:\n- Client Devices: PCs/Macs and mobile devices to test Responsive UI built with Tailwind CSS v4.\n- Server Environment: Vercel Cloud infrastructure for Next.js 16 deployments. Local/Cloud Python servers for the FastAPI engine requiring sufficient RAM for AI orchestration.\n- Network: Continuous high-speed access for Firebase Storage, Pinecone Vector DB, and external AI APIs (OpenRouter, ElevenLabs, OpenAI)."

    if "[Note:  Add or delete items as appropriate.]" in p.text:
        p.text = ""
    if "[Note:  Add more items as appropriate.]" in p.text:
        p.text = ""

if len(doc.tables) > 0:
    rev_table = doc.tables[0]
    if len(rev_table.rows) > 1:
        rev_table.rows[1].cells[0].text = "24/Jun/2026"
        rev_table.rows[1].cells[1].text = "3.0"
        rev_table.rows[1].cells[2].text = "Updated Test Plan reflecting PA3 SAD-v3 Architecture and Working Software"
        rev_table.rows[1].cells[3].text = "Vũ, Hoàng, Hùng"

if len(doc.tables) > 2:
    prod_table = doc.tables[2]
    if len(prod_table.rows) > 1:
        prod_table.rows[1].cells[0].text = "Issue & Bug Tracking"
        prod_table.rows[1].cells[1].text = "Jira"
        prod_table.rows[1].cells[2].text = "Atlassian"
        prod_table.rows[1].cells[3].text = "Cloud"
    if len(prod_table.rows) > 2:
        prod_table.rows[2].cells[0].text = "Version Control"
        prod_table.rows[2].cells[1].text = "GitHub"
        prod_table.rows[2].cells[2].text = "GitHub"
        prod_table.rows[2].cells[3].text = "Cloud"

if len(doc.tables) > 3:
    roles_table = doc.tables[3]
    if len(roles_table.rows) > 1:
        roles_table.rows[1].cells[1].text = "Hoàng (1)"
    if len(roles_table.rows) > 2:
        roles_table.rows[2].cells[1].text = "Vũ, Hùng (2)"
    if len(roles_table.rows) > 3:
        roles_table.rows[3].cells[1].text = "Hoàng (1)"

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_p(p, "<Project Name>", "Melon Learning App")
                replace_in_p(p, "<1.0>", "3.0")

try:
    doc.save(dest_doc)
    if os.path.exists(final_doc):
        os.remove(final_doc)
except Exception as e:
    print(e)
