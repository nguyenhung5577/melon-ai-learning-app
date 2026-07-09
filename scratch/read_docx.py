from docx import Document
doc = Document(r'd:\HCMUS\HK2-N3\SE4AI\melon-ai-learning-app\rup_tstpln.docx')
for p in doc.paragraphs:
    print(p.text)
for t in doc.tables:
    for row in t.rows:
        print(' | '.join([c.text.replace('\n', ' ') for c in row.cells]))
